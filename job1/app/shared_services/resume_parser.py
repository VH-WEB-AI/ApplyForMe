"""
Shared Resume Parser service (Section 5: Shared AI Services -> RESUME PARSER).
Extracts raw text from PDF/DOCX/TXT resumes. Used by Engine 1 (Resume
Intelligence) and by the Career Copilot's context loader.
"""
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from app.core.exceptions import ValidationFailedError

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class ResumeParser:
    def parse(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValidationFailedError(f"Unsupported resume file type: {ext}")

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        if ext == ".docx":
            return self._parse_docx(file_path)
        return self._parse_txt(file_path)

    def _parse_pdf(self, file_path: str) -> str:
        """Try PyMuPDF first (fast, good layout handling); fall back to
        pdfplumber for trickier PDFs (tables, unusual encodings)."""
        try:
            text_parts = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    text_parts.append(page.get_text())
            text = "\n".join(text_parts).strip()
            if text:
                return text
        except Exception:
            pass

        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()

    def _parse_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()

    def _parse_txt(self, file_path: str) -> str:
        return Path(file_path).read_text(encoding="utf-8", errors="ignore").strip()


resume_parser = ResumeParser()
