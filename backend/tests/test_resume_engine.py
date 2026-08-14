import io
from unittest.mock import patch

import docx

from app.db.models.core import CandidateProfile, User
from app.db.models.resume import ResumeVersion
from app.engines.resume_intelligence.engine import ResumeIntelligenceEngine
from app.engines.resume_intelligence.schemas import ResumeLLMOutput
from app.orchestrator.orchestrator import AIOrchestrator
from app.services.llm_gateway import ChatResult, EmbeddingResult

FAKE_EMBEDDING = EmbeddingResult(vector=[0.01] * 1536, model="test-embedding-model", latency_ms=1.0)

FAKE_LLM_OUTPUT = ResumeLLMOutput(
    tags=["python", "fastapi", "kubernetes"],
    resume_score=72,
    ats_score=68,
    section_scores={"summary": 80, "experience": 65, "education": 80, "skills": 60},
    weak_sections=["experience"],
    total_experience_years=3.0,
    education=["Bachelor of Science in Computer Science, State University, 2018"],
    certifications=["AWS Certified Solutions Architect"],
    missing_skills=["Terraform"],
    recommendations=["Quantify the 'worked on stuff' bullet with a measurable outcome."],
    rewrite_suggestions={"experience": "Owned X, resulting in Y% improvement in Z."},
)


def _make_resume_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane.doe@example.com | 555-123-4567")
    document.add_paragraph("Summary")
    document.add_paragraph("Backend engineer with a focus on distributed systems and APIs.")
    document.add_paragraph("Experience")
    document.add_paragraph("- Led migration of 12 services to Kubernetes, reducing costs by 30%")
    document.add_paragraph("- Built a FastAPI platform used by 200 engineers")
    document.add_paragraph("- worked on stuff")
    document.add_paragraph("Education")
    document.add_paragraph("Bachelor of Science in Computer Science, State University, 2018")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL, Docker, Kubernetes, AWS")
    document.add_paragraph("Certifications")
    document.add_paragraph("AWS Certified Solutions Architect")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_candidate(db) -> CandidateProfile:
    user = User(email="jane@example.com", full_name="Jane Doe")
    db.add(user)
    db.flush()
    profile = CandidateProfile(user_id=user.id, target_role="Backend Engineer")
    db.add(profile)
    db.flush()
    return profile


def test_gather_context_has_no_pii_and_defers_tags_to_llm(db):
    candidate = _make_candidate(db)
    engine = ResumeIntelligenceEngine()

    with patch("app.services.embedding_generator.llm_gateway.create_embedding", return_value=FAKE_EMBEDDING):
        context = engine.gather_context(
            db,
            {
                "candidate_id": candidate.id,
                "file_bytes": _make_resume_docx(),
                "filename": "resume.docx",
                "target_role": "Backend Engineer",
                "target_industry": "Software",
            },
        )

    assert context["resume_version_id"] is not None
    # PII must not leak into what eventually reaches the LLM
    assert "jane.doe@example.com" not in context["redacted_resume_text"]
    # tags aren't known until the LLM call in postprocess() runs
    resume_version = db.get(ResumeVersion, context["resume_version_id"])
    assert resume_version.tags == []


def test_resume_engine_full_orchestrator_flow(db):
    candidate = _make_candidate(db)

    fake_chat_result = ChatResult(
        content=FAKE_LLM_OUTPUT.model_dump_json(),
        model="test-model",
        prompt_tokens=100,
        completion_tokens=50,
        latency_ms=250.0,
    )

    payload = {
        "candidate_id": candidate.id,
        "file_bytes": _make_resume_docx(),
        "filename": "resume.docx",
        "target_role": "Backend Engineer",
        "target_industry": "Software",
    }

    with (
        patch("app.orchestrator.orchestrator.chat_completion", return_value=fake_chat_result),
        patch("app.services.embedding_generator.llm_gateway.create_embedding", return_value=FAKE_EMBEDDING),
    ):
        result = AIOrchestrator().handle_request("resume_intelligence", db, payload)

    assert result["resumeScore"] == 72
    assert result["atsScore"] == 68
    assert result["tags"] == ["python", "fastapi", "kubernetes"]
    assert result["sectionScores"]["experience"] == 65
    assert result["weakSections"] == ["experience"]
    assert result["totalExperienceYears"] == 3.0
    assert any("AWS Certified" in c for c in result["certifications"])
    assert "Bachelor" in " ".join(result["education"])
    assert result["missingSkills"] == ["Terraform"]
    assert "recommendations" in result
    assert result["rewriteSuggestions"]["experience"].startswith("Owned")

    # tags are persisted back onto the resume version row once the LLM call returns
    resume_version = db.get(ResumeVersion, result["resumeVersionId"])
    assert resume_version.tags == ["python", "fastapi", "kubernetes"]


def test_ats_check_without_candidate_id_still_returns_full_analysis(db):
    fake_chat_result = ChatResult(
        content=FAKE_LLM_OUTPUT.model_dump_json(),
        model="test-model",
        prompt_tokens=100,
        completion_tokens=50,
        latency_ms=250.0,
    )

    payload = {
        "candidate_id": None,
        "file_bytes": _make_resume_docx(),
        "filename": "resume.docx",
        "target_role": "Backend Engineer",
        "target_industry": "Software",
    }

    with patch("app.orchestrator.orchestrator.chat_completion", return_value=fake_chat_result):
        result = AIOrchestrator().handle_request("resume_intelligence", db, payload)

    assert result["resumeVersionId"] is None
    assert result["tags"] == ["python", "fastapi", "kubernetes"]
    assert result["resumeScore"] == 72
